import re
import time
import openai
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
import os, fitz, spacy
from flask_socketio import SocketIO

ner_models = {}
load_dotenv('config.env')
API_KEY = os.getenv('API_KEY')
client = OpenAI(api_key=API_KEY)

def split_into_paragraphs(text):
    return [paragraph for paragraph in text.split('\n') if paragraph.strip() != '']

def split_into_sentences(paragraph):
    return paragraph.split('. ')

def get_docx_text(file_path):
    doc = Document(file_path)
    whole_text = ""
    for paragraph in doc.paragraphs:
        whole_text += paragraph.text + "\n"
    return whole_text
def get_pdf_paragraphs(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        page_text = page.get_text("text")
        paragraphs = page_text.split('\n')
        for paragraph in paragraphs:
            text += paragraph + "\n"
    doc.close()
    return text

def remove_footnote_markers(text):
    pattern = r'\[\d+\]|\d+\s?|\(\d+\)' 
    cleaned_text = re.sub(pattern, '', text)
    return cleaned_text

def translate_text(sentences, client, custom_prompt):
    context_messages = [{"role": "system", "content": f"{custom_prompt}"}]
    for sentence in sentences[-4:]:  # Keep the context of the last 3 sentences plus the current one
        context_messages.append({"role": "user", "content": sentence})

    while True:
        try:
            response = client.chat.completions.create(model="gpt-4",
                                                      messages=context_messages,
                                                      temperature=0,
                                                      max_tokens=6969)
            return response.choices[0].message.content.strip()

        except openai.RateLimitError:
            time.sleep(65)
        except openai.OpenAIError as e:
            raise e

def create_translated_document(paragraphs, client, output_file_path, socketio, custom_prompt):
    translated_doc = Document()

    for index, (paragraph, style) in enumerate(paragraphs):
        sentences = split_into_sentences(paragraph)
        translated_paragraph = []
        for i in range(len(sentences)):
            context_sentences = sentences[max(0, i-3):i+1] 
            translated_sentence = translate_text(context_sentences,  client, custom_prompt)
            translated_paragraph.append(translated_sentence.split('\n')[-1]) 
        
        progress = (index + 1) / len(paragraphs) * 100 
        socketio.emit('progress_update', {'progress': progress})
        p = translated_doc.add_paragraph(' '.join(translated_paragraph))
        p.style = style

    translated_doc.save(output_file_path)

def ner_translate(entity, context_window, source_language, target_language, client):
    print(f"translating {entity}")
    context_messages = [{"role": "system", "content": f"If you cannot translate {entity}, DO NOT leave comments. Translate this text '{entity}' from {source_language} to {target_language} academically, given this context: '{context_window}' translate this text '{entity}'."}]

    while True:
        try:
            response = client.chat.completions.create(model="gpt-4",
                                                      messages=context_messages,
                                                      temperature=0.4,
                                                      max_tokens=6969)
            return response.choices[0].message.content.strip()

        except openai.RateLimitError:
            time.sleep(65)
        except openai.OpenAIError as e:
            raise e

def identify_entities(text, source_language, target_language):

    if source_language == "en":
        ner_models['en'] = spacy.load('en' + '_core_web_sm')
    if (source_language == "zh-cn"):
        source_language = "zh"
        ner_models[source_language] = spacy.load(f'{source_language}_core_web_sm')
    if (source_language == "ru"):
        ner_models[source_language] = spacy.load(f'{source_language}_core_web_sm')
        
    nlp = ner_models[source_language]
    doc = nlp(text)
    entities = {}

    for ent in doc.ents:
        if ent.text not in entities and ent.label_ in ['PERSON', 'GPE', 'LOC', 'ORG', 'FAC', 'EVENT', 'NORP', 'WORK_OF_ART', 'PRODUCT']:
            full_sentence = ent.sent.text
            entities[ent.text] = ner_translate(ent.text, full_sentence, source_language=source_language, target_language=target_language, client=client)
            print(entities[ent.text])
    return entities

def replace_entities(text, mapping):
    regex_pattern = '|'.join(re.escape(key) for key in mapping.keys())
    def replace_match(match):
        return mapping[match.group(0)]
    modified_text = re.sub(regex_pattern, replace_match, text)
    return modified_text